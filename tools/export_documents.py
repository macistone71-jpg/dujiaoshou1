#!/usr/bin/env python3
"""Export the six website articles as five WeChat-style Word documents.

Requires: python-docx. Article JSON and PNGs are prepared by the export command
listed in README.md.
"""
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/tmp/dujiaoshou-articles.json")
PNG_DIR = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("/tmp/dujiaoshou-png")
OUT = ROOT / "交付文档"
OUT.mkdir(exist_ok=True)

PACKAGES = [
    ("01-产品价值与需求分析（含2篇）.docx", [0, 1], "产品价值与需求分析"),
    ("02-产品经典阅读与认知升级.docx", [2], "产品经典阅读与认知升级"),
    ("03-做减法：为什么好产品都是克制的.docx", [3], "做减法：为什么好产品都是克制的"),
    ("04-用AI做产品：实践、工作流与边界.docx", [4], "用 AI 做产品：实践、工作流与边界"),
    ("05-从执行到负责：一个小功能上线复盘.docx", [5], "从执行到负责：一个小功能上线复盘"),
]


def set_east_asia(run, font="PingFang SC"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_shading(paragraph, fill="F7F7F7"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_left_border(paragraph, color="D26911", size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)

    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(63, 63, 63)
    normal.paragraph_format.line_spacing = 1.75
    normal.paragraph_format.space_after = Pt(10)

    for name, size, bold, color, before, after in [
        ("公众号标题", 22, True, "1F1F1F", 0, 10),
        ("公众号二级标题", 15, True, "202020", 22, 10),
        ("公众号元信息", 9.5, False, "999999", 0, 16),
        ("公众号图片说明", 9, False, "999999", 3, 16),
        ("公众号导语", 10.5, False, "666666", 0, 18),
    ]:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.45 if "标题" in name else 1.6


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph(caption, style="公众号图片说明")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_block(doc, block):
    if block.get("h2"):
        p = doc.add_paragraph(block["h2"], style="公众号二级标题")
        set_left_border(p)
        return
    if block.get("p"):
        p = doc.add_paragraph(block["p"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return
    if block.get("quote"):
        p = doc.add_paragraph(block["quote"])
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.right_indent = Cm(0.2)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(14)
        set_cell_shading(p, "FAF6F2")
        set_left_border(p)
        return
    if block.get("list"):
        for item in block["list"]:
            p = doc.add_paragraph(item, style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.2)
            p.paragraph_format.line_spacing = 1.6
            p.paragraph_format.space_after = Pt(5)


def png_for(visual):
    return PNG_DIR / (Path(visual["src"]).stem + ".png")


def add_article(doc, post, visuals, with_page_break=False):
    if with_page_break:
        doc.add_page_break()

    label = doc.add_paragraph(f'{post["category"]} · 产品思考')
    label.style = doc.styles["公众号元信息"]
    if label.runs:
        label.runs[0].font.color.rgb = RGBColor(210, 105, 17)
        label.runs[0].font.bold = True

    title = doc.add_paragraph(post["title"], style="公众号标题")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    chars = sum(len(str(b.get("p") or b.get("h2") or b.get("quote") or "")) + sum(len(x) for x in b.get("list", [])) for b in post["content"])
    meta = doc.add_paragraph(f'何庆丰    {post["date"]}    约 {max(1, (chars + 449) // 450)} 分钟阅读', style="公众号元信息")

    add_picture(doc, png_for(visuals[0]), visuals[0]["caption"])

    lead = doc.add_paragraph(post["excerpt"], style="公众号导语")
    lead.paragraph_format.left_indent = Cm(0.35)
    lead.paragraph_format.right_indent = Cm(0.35)
    set_cell_shading(lead)

    blocks = post["content"]
    inline = visuals[1:]
    slots = {}
    for idx, visual in enumerate(inline):
        slot = max(1, round(((idx + 1) / (len(inline) + 1)) * len(blocks)))
        while slot in slots and slot < len(blocks) - 1:
            slot += 1
        slots[slot] = visual

    for index, block in enumerate(blocks, 1):
        add_block(doc, block)
        if index in slots:
            visual = slots[index]
            add_picture(doc, png_for(visual), visual["caption"])

    end = doc.add_paragraph("END\n何庆丰\n日拱一卒，功不唐捐")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.paragraph_format.space_before = Pt(28)
    end.paragraph_format.line_spacing = 1.5
    for run in end.runs:
        set_east_asia(run)
        run.font.color.rgb = RGBColor(130, 130, 130)


def build_document(filename, indices, package_title, posts, visuals):
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = package_title
    doc.core_properties.author = "何庆丰"
    doc.core_properties.subject = "公众号格式文章排版稿"
    doc.core_properties.comments = "由个人网站内容导出；每篇文章包含 5 张原创插图。"

    for n, index in enumerate(indices):
        add_article(doc, posts[index], visuals[index], with_page_break=n > 0)

    path = OUT / filename
    doc.save(path)
    return path


def main():
    payload = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    posts, visuals = payload["posts"], payload["visuals"]
    paths = [build_document(*package, posts, visuals) for package in PACKAGES]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
